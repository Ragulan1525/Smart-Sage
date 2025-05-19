# import chromadb
# from sentence_transformers import SentenceTransformer
# import wikipedia
# import requests
# import os
# import logging
# import asyncio
# from dotenv import load_dotenv
# from bs4 import BeautifulSoup
# from Backend.storage import get_chroma_collection
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from newspaper import Article
# import PyPDF2
# import docx
# from docx import Document
# import fitz
# import httpx
# import uuid
# import validators
# import re
# import nltk
# from nltk.tokenize import sent_tokenize
# from transformers import pipeline
# from pdfminer.high_level import extract_text as extract_text_from_pdf
# from functools import lru_cache
# from concurrent.futures import ThreadPoolExecutor
# import time
# from typing import List, Dict
# from datetime import datetime
# from Backend.shared_utils import process_search_results

# # Load environment variables from .env file
# load_dotenv()
# GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
# GOOGLE_SEARCH_ENGINE_ID = os.getenv("GOOGLE_SEARCH_ENGINE_ID")
# DUCKDUCKGO_API_URL = "https://api.duckduckgo.com/"
# NEWS_API_KEY = os.getenv("NEWS_API_KEY")

# logging.basicConfig(level=logging.DEBUG)

# # Load embedding model
# model = SentenceTransformer('all-MiniLM-L6-v2')
# logging.info("✅ Model loaded successfully!")

# # Load ChromaDB collection once
# chroma_collection = get_chroma_collection()


# # Initialize ChromaDB
# chroma_client = chromadb.PersistentClient(path="./chroma_db")
# collection = chroma_client.get_or_create_collection(name="study_docs")


# nltk.download("punkt")

# # Educational websites for scraping
# EDUCATIONAL_SITES = [
#     "https://www.khanacademy.org/",
#     "https://www.coursera.org/",
#     "https://www.edx.org/",
#     "https://www.geeksforgeeks.org/"
# ]


# # ✅ Load Summarization Model (One-time)
# summarizer = pipeline("summarization", model="facebook/bart-large-cnn")


# async def search_google(query: str, model, chroma_collection) -> List[str]:
#     """Fetch live search results from Google, fallback to DuckDuckGo if Google fails."""
#     if not query or len(query.strip()) < 3:
#         logging.warning("Query too short for Google search")
#         return []

#     url = "https://www.googleapis.com/customsearch/v1"
#     params = {
#         "q": query,
#         "key": GOOGLE_API_KEY,
#         "cx": GOOGLE_SEARCH_ENGINE_ID,
#         "num": 3,  # Number of results to fetch
#         "dateRestrict": "d1"
#     }

#     retries = 3  # Number of retry attempts
#     last_error = None

#     async with httpx.AsyncClient(timeout=10.0) as client:
#         for attempt in range(retries):
#             try:
#                 response = await client.get(url, params=params, timeout=10)
#                 response.raise_for_status()

#                 if response.status_code == 200:
#                     data = response.json()
#                     if "items" in data and data["items"]:
#                         results = await process_search_results(data, query, model, chroma_collection)
#                         if results:
#                             return results
#                     else:
#                         logging.warning("⚠️ No results from Google, trying DuckDuckGo...")
#                         return await search_duckduckgo(query, model, chroma_collection)

#             except httpx.RequestError as e:
#                 last_error = e
#                 wait_time = 2 ** attempt  # Exponential backoff
#                 logging.warning(f"⚠️ Network error: {e}, Retrying in {wait_time} seconds...")
#                 await asyncio.sleep(wait_time)
#                 continue

#             except httpx.HTTPStatusError as e:
#                 if e.response.status_code == 429:
#                     logging.warning("⚠️ Google API Rate Limit hit. Switching to DuckDuckGo.")
#                     return await search_duckduckgo(query, model, chroma_collection)
#                 else:
#                     logging.error(f"❌ Google API Error {e.response.status_code}: {e.response.text}")
#                     last_error = e
#                     break

#     logging.warning(f"⚠️ Google search failed after {retries} attempts: {last_error}")
#     return await search_duckduckgo(query, model, chroma_collection)


# async def search_duckduckgo(query, model, chroma_collection):
#     """Fetches search results from DuckDuckGo if Google fails."""
#     try:
#         async with httpx.AsyncClient() as client:
#             params = {"q": query, "format": "json"}
#             response = await client.get(DUCKDUCKGO_API_URL, params=params, timeout=10)
#             response.raise_for_status()
#             data = response.json()

#             related_topics = data.get("RelatedTopics", [])
#             extracted_results = []

#             for topic in related_topics:
#                 if "Text" in topic and "FirstURL" in topic:
#                     url = topic["FirstURL"]
#                     if not url.startswith("http"):
#                         continue
#                     extracted_results.append(topic["Text"])

#             if extracted_results:
#                 await store_text_in_chroma(
#                     "\n".join(extracted_results),
#                     f"DuckDuckGo Data ({query})",
#                     model,
#                     chroma_collection
#                 )

#             if not extracted_results:
#                 logging.warning(f"⚠️ No useful DuckDuckGo results found for '{query}'")

#             return extracted_results

#     except Exception as e:
#         logging.error(f"❌ Unexpected error in DuckDuckGo search: {e}", exc_info=True)
#         return []


# async def fetch_latest_news(query):
#     """Fetches the latest news from NewsAPI."""
#     try:
#         if not NEWS_API_KEY:
#             raise ValueError("❌ Missing NewsAPI Key!")

#         url = f"https://newsapi.org/v2/everything?q={query}&apiKey={NEWS_API_KEY}&language=en&pageSize=5"

#         async with httpx.AsyncClient() as client:
#             response = await client.get(url, timeout=10)
#             response.raise_for_status()
#             data = response.json()

#         return [f"{article['title']}: {article['url']}" for article in data.get("articles", [])] if data.get("articles") else None
#     except Exception as e:
#         logging.error(f"❌ News API Error: {e}")
#         return None




# import httpx
# import logging
# from bs4 import BeautifulSoup
# from urllib.parse import urlparse
# from transformers import pipeline

# # Initialize only once (ensure you do this outside if you're importing it)
# summarizer = pipeline("summarization")

# def is_valid_url(url: str) -> bool:
#     """Validates whether a string is a proper URL."""
#     try:
#         parsed = urlparse(url)
#         return parsed.scheme in ("http", "https") and parsed.netloc != ""
#     except Exception:
#         return False

# # Add caching for expensive operations
# @lru_cache(maxsize=1000)
# def get_embedding(text: str) -> List[float]:
#     """Cache embeddings for frequently used text."""
#     return model.encode(text).tolist()

# # Add parallel processing for multiple sources
# async def parallel_retrieve(query: str, model, chroma_collection) -> Dict[str, List[str]]:
#     """Retrieve information from multiple sources in parallel with optimized timeouts."""
#     try:
#         # Validate query
#         if not query or len(query.strip()) < 3:
#             logging.warning("Query too short for retrieval")
#             return {"error": "Query too short"}

#         # Set timeouts for each source with different values based on source reliability
#         timeouts = {
#             "chroma": 15.0,  # Match with app.py
#             "wikipedia": 10.0,
#             "google": 20.0,
#             "news": 10.0
#         }
        
#         # Create tasks with individual timeouts
#         tasks = {
#             "chroma": asyncio.wait_for(retrieve_relevant_text(query, model, chroma_collection), timeout=timeouts["chroma"]),
#             "wikipedia": asyncio.wait_for(async_wikipedia_search(query), timeout=timeouts["wikipedia"]),
#             "google": asyncio.wait_for(search_google(query, model, chroma_collection), timeout=timeouts["google"]),
#             "news": asyncio.wait_for(fetch_latest_news(query), timeout=timeouts["news"])
#         }
        
#         # Execute all tasks in parallel with return_exceptions=True
#         results = await asyncio.gather(*tasks.values(), return_exceptions=True)
        
#         # Process results with detailed error logging
#         processed_results = {}
#         for source, result in zip(tasks.keys(), results):
#             if isinstance(result, asyncio.TimeoutError):
#                 logging.warning(f"Timeout while retrieving from {source} after {timeouts[source]} seconds")
#                 processed_results[source] = []
#             elif isinstance(result, Exception):
#                 logging.warning(f"Error retrieving from {source}: {str(result)}")
#                 processed_results[source] = []
#             elif result is None:
#                 logging.info(f"No results from {source}")
#                 processed_results[source] = []
#             else:
#                 # Clean and validate the results
#                 if isinstance(result, list):
#                     processed_results[source] = [r for r in result if r and isinstance(r, str)]
#                 else:
#                     processed_results[source] = [str(result)] if result else []
                
#                 if not processed_results[source]:
#                     logging.info(f"Empty results after cleaning from {source}")

#         # Log the results for debugging
#         for source, results in processed_results.items():
#             logging.info(f"Results from {source}: {len(results)} items")
#             if results:
#                 logging.debug(f"First result from {source}: {results[0][:100]}...")

#         # If all sources failed, try a simpler approach with ChromaDB
#         if all(not results for results in processed_results.values()):
#             return {"error": "All sources failed. Please try again later or refine your query."}
#             logging.warning("All sources failed, trying ChromaDB with extended timeout")
#             try:
#                 chroma_result = await asyncio.wait_for(
#                     retrieve_relevant_text(query, model, chroma_collection),
#                     timeout=15.0  # Extended timeout for fallback
#                 )
#                 processed_results["chroma"] = chroma_result if chroma_result else []
#             except Exception as e:
#                 logging.error(f"ChromaDB fallback failed: {e}")
#                 processed_results["chroma"] = []
                
#         return processed_results
        
#     except Exception as e:
#         logging.error(f"Critical error in parallel retrieval: {e}")
#         # Try to get at least ChromaDB results with extended timeout
#         try:
#             chroma_result = await asyncio.wait_for(
#                 retrieve_relevant_text(query, model, chroma_collection),
#                 timeout=15.0
#             )
#             return {"chroma": chroma_result if chroma_result else []}
#         except Exception as e:
#             logging.error(f"Final ChromaDB fallback failed: {e}")
#             return {"error": "Failed to retrieve information from any source"}

# # Optimize ChromaDB queries
# async def retrieve_relevant_text(query: str, model, top_k: int = 3) -> List[str]:
#     """Optimized retrieval from ChromaDB with caching."""
#     try:
#         # Get embedding with caching
#         query_embedding = get_embedding(query)
        
#         # Query ChromaDB with optimized parameters.count()
#         results = chroma_collection.query(
#             query_embeddings=[query_embedding], int(available_docs)
#             n_results=top_k,
#             include=["documents", "metadatas", "distances"] Query ChromaDB with optimized parameters
#         )num_results = max(1, min(top_k, available_docs))
        
#         if not results or not results["documents"]:
#             logging.warning(f"No relevant documents found for query: {query}")
#             return "No relevant results found."include=["documents", "metadatas", "distances"]
            
#         # Process results with metadata and distance filtering
#         processed_results = []
#         for doc, meta, distance in zip(results["documents"][0], results["metadatas"][0], results["distances"][0]):
#             if doc and distance < 0.3:  # Only include results with good similarity
#                 processed_results.append(f"{doc} (Source: {meta.get('source', 'Unknown')})")    
        
#         return processed_results if processed_results else "No relevant results found."processed_results = []
#         tance in zip(results["documents"][0], results["metadatas"][0], results["distances"][0]):
#     except Exception as e:larity
#         logging.error(f"Error in retrieve_relevant_text: {e}", exc_info=True){meta.get('source', 'Unknown')})")
#         return f"Error retrieving from ChromaDB: {str(e)}"        
# d_results if processed_results else "No relevant results found."
# # Optimize web scraping
# async def scrape_and_summarize(url_or_query: str) -> str:
#     """Optimized web scraping with caching and timeout."""logging.error(f"Error in retrieve_relevant_text: {e}", exc_info=True)
#     try:ieving from ChromaDB: {str(e)}"
#         # Check cache first
#         cache_key = f"scrape_{hash(url_or_query)}"
#         cached = chroma_collection.get(ids=[cache_key])r_query: str) -> str:
#         if cached and cached["ids"]:g and timeout."""
#             return cached["documents"][0]
            
#         # Set timeout for scrapingscrape_{hash(url_or_query)}"
#         timeout = 10.0on.get(ids=[cache_key])
#         start_time = time.time()if cached and cached["ids"]:
        
#         async with httpx.AsyncClient(timeout=timeout) as client:
#             if validators.url(url_or_query):
#                 response = await client.get(url_or_query)
#                 text = response.texte = time.time()
#             else:
#                 # Handle query-based scraping as client:
#                 text = await search_web(url_or_query)alidators.url(url_or_query):
#                 _query)
#             if time.time() - start_time > timeout:
#                 raise TimeoutError("Scraping timed out"):
#                 d scraping
#             # Extract and summarize
#             soup = BeautifulSoup(text, 'html.parser')
#             content = soup.get_text()
#             summary = extract_first_sentences(content)    raise TimeoutError("Scraping timed out")
            
#             # Cache the result
#             await store_text_in_chroma(summary, "Web Scrape", model, cache_key)soup = BeautifulSoup(text, 'html.parser')
#             .get_text()
#             return summarysummary = extract_first_sentences(content)
            
#     except Exception as e:
#         logging.error(f"Error in web scraping: {e}") store_text_in_chroma(summary, "Web Scrape", model, cache_key)
#         return ""            

# # ✅ Helper Function to Extract First Full Sentences (Avoid Mid-Cutoff)
# def extract_first_sentences(text, max_length=500):
#     """Extracts the first few full sentences without cutting off mid-way.""" web scraping: {e}")
#     sentences = text.split(". ")""
#     output = ""
#     for sentence in sentences:(Avoid Mid-Cutoff)
#         if len(output) + len(sentence) > max_length:_sentences(text, max_length=500):
#             breakl sentences without cutting off mid-way."""
#         output += sentence + ". "t(". ")
#     return output.strip()    output = ""
#     for sentence in sentences:
#  len(sentence) > max_length:
# # def search_web(query):
# #     """Performs a web search using DuckDuckGo API."""tput += sentence + ". "
# #     try:
# #         url = f"https://api.duckduckgo.com/?q={query}&format=json"
# #         response = requests.get(url, timeout=10)
# #         response.raise_for_status()
# #         data = response.json()
# #         results = [topic["Text"] for topic in data.get("RelatedTopics", []) if "Text" in topic]
# #         return results if results else ["No relevant results found."]&format=json"
# #     except requests.exceptions.RequestException as e:
# #         logging.error(f"Web search failed: {e}")
# #         return ["Web search error."]on()
# #     except Exception as e:latedTopics", []) if "Text" in topic]
# #         logging.error(f"Unexpected web search error: {e}")und."]
# #         return ["Unexpected error occurred during web search."]#     except requests.exceptions.RequestException as e:
# ng.error(f"Web search failed: {e}")
# import requests#         return ["Web search error."]
# #     except Exception as e:
# r(f"Unexpected web search error: {e}")
# def search_web(query):eb search."]
#     """Performs a web search using DuckDuckGo API."""
#     try:
#         url = f"https://api.duckduckgo.com/?q={query}&format=json"
#         response = requests.get(url, timeout=10)
#         response.raise_for_status()
#         data = response.json()    """Performs a web search using DuckDuckGo API."""

#         results = [topic["Text"] for topic in data.get("RelatedTopics", []) if "Text" in topic]&format=json"
#     except requests.exceptions.RequestException as e:
#         logging.error(f"Web search failed: {e}")
#         return ["Web search error."]        data = response.json()

# et("RelatedTopics", []) if "Text" in topic]
# async def scrape_educational_websites(query, model):
#     """Scrapes educational websites for study material using semantic similarity."""error(f"Web search failed: {e}")
#     results = []
#     query_embedding = model.encode(query)

#     for site in EDUCATIONAL_SITES:rape_educational_websites(query, model):
#         try:l using semantic similarity."""
#             async with httpx.AsyncClient() as client:
#                 response = await client.get(site, timeout=10)
#                 response.raise_for_status()

#             soup = BeautifulSoup(response.text, "html.parser")
#             content_elements = soup.find_all(['p', 'div', 'article', 'section', 'span'])
#             text = " ".join([element.get_text(separator=" ", strip=True) for element in content_elements])                response = await client.get(site, timeout=10)

#             text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
#             chunks = text_splitter.split_text(text)            soup = BeautifulSoup(response.text, "html.parser")
# oup.find_all(['p', 'div', 'article', 'section', 'span'])
#             for chunk in chunks:or=" ", strip=True) for element in content_elements])
#                 chunk_embedding = model.encode(chunk)
#                 similarity = chunk_embedding @ query_embeddingeCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
#                 if similarity > 0.6:
#                     results.append((similarity, f"From {site}: {chunk}"))
# s:
#         except Exception as e:
#             logging.error(f"Error fetching or processing {site}: {e}")                similarity = chunk_embedding @ query_embedding

#     results.sort(key=lambda x: x[0], reverse=True)f"From {site}: {chunk}"))
#     return [result[1] for result in results[:3]]
#         except Exception as e:
# ing.error(f"Error fetching or processing {site}: {e}")
# import wikipedia
# erse=True)
# async def async_wikipedia_search(query):
#     """Fetch Wikipedia summary with proper error handling."""
#     try:
#         search_results = wikipedia.search(query)
#         if not search_results:
#             logging.warning(f"⚠️ No Wikipedia results for query: {query}")ia_search(query):
#             return None    """Fetch Wikipedia summary with proper error handling."""

#         for result in search_results[:3]:esults = wikipedia.search(query)
#             try:
#                 page = wikipedia.page(result, auto_suggest=True)ery}")
#                 summary = wikipedia.summary(page.title, sentences=3)
#                 if len(summary) > 50:
#                     return summary
#             except wikipedia.exceptions.DisambiguationError as e:
#                 logging.warning(f"⚠️ Disambiguation error for {result}: {e}")ikipedia.page(result, auto_suggest=True)
#                 continueitle, sentences=3)
#             except wikipedia.exceptions.PageError:
#                 logging.warning(f"⚠️ Page error for {result}")rn summary
#                 continue            except wikipedia.exceptions.DisambiguationError as e:
# ging.warning(f"⚠️ Disambiguation error for {result}: {e}")
#         return None
#     except Exception as e:
#         logging.error(f"❌ Wikipedia API error: {e}")ging.warning(f"⚠️ Page error for {result}")
#         return None                continue

#             return None

# async def async_scrape_web(query):
#     """Asynchronously scrape educational websites for content if needed."""
#     return scrape_educational_websites(query, model) or None
    


# async def retrieve_relevant_text(query: str, model, top_k=3):"""
#     """Retrieve relevant study material asynchronously from ChromaDB."""rn scrape_educational_websites(query, model) or None
#     try:
#         query_embedding = model.encode(query).tolist()
#         available_docs = chroma_collection.count()
#         num_results = max(1, min(top_k, available_docs))async def retrieve_relevant_text(query: str, model, top_k=3):
# ynchronously from ChromaDB."""
#         results = await asyncio.to_thread(
#             chroma_collection.query, list()
#             query_embeddings=[query_embedding], ollection.count()  # Ensure this returns an integer
#             n_results=num_resultsf not isinstance(available_docs, int):
#         )            available_docs = int(available_docs)  # Convert to integer if necessary

#         if not results or "documents" not in results or not results["documents"]:arison
#             logging.warning(f"No relevant documents found for query: {query}")
#             return "No relevant results found."        results = await asyncio.to_thread(

#         retrieved_docs = results["documents"][0] if results["documents"] else []
#         distances = results.get("distances", [[]])[0] if results.get("distances") else []            n_results=num_results

#         cleaned_results = []
#         for idx, text in enumerate(retrieved_docs):results or not results["documents"]:
#             if text and distances[idx] < 0.3:
#                 text = re.sub(r"[\n\t]+", " ", text)  # Remove line breaks/tabs
#                 text = re.sub(r"http\S+", "", text)  # Remove links
#                 text = re.sub(r"[^a-zA-Z0-9.,!?'\- ]+", "", text)  # ✅ Keeps hyphens and apostropheslts["documents"] else []
#                 text = re.sub(r"\s+", " ", text).strip()s = results.get("distances", [[]])[0] if results.get("distances") else []
                
#                 if len(text.split()) > 5:
#                     cleaned_results.append(text)        for idx, text in enumerate(retrieved_docs):

#         return "\n\n".join(cleaned_results) if cleaned_results else "No relevant results found."                text = re.sub(r"[\n\t]+", " ", text)  # Remove line breaks/tabs
# sub(r"http\S+", "", text)  # Remove links
#     except Exception as e:hyphens and apostrophes
#         logging.error(f"Error in retrieve_relevant_text: {e}", exc_info=True)
#         return f"Error retrieving from ChromaDB: {str(e)}"            
#                     if len(text.split()) > 5:
#                     cleaned_results.append(text)

# import fitz  # PyMuPDF        return "\n\n".join(cleaned_results) if cleaned_results else "No relevant results found."

# def extract_text_from_pdf(file_obj):
#     """Attempts text extraction from PDF using PyMuPDF, with pdfminer fallback."""logging.error(f"Error in retrieve_relevant_text: {e}", exc_info=True)
#     try:
#         with fitz.open(stream=file_obj.file.read(), filetype="pdf") as doc:
#             text = ""
#             for page in doc:
#                 text += page.get_text()
#         if not text.strip():
#             raise ValueError("Empty text extracted with PyMuPDF.")om_pdf(file_obj):
#         return textction from PDF using PyMuPDF, with pdfminer fallback."""
#     except Exception as e:
#         logging.warning(f"⚠️ PyMuPDF failed: {e}. Trying fallback with pdfminer.")ad(), filetype="pdf") as doc:
#         # Rewind file pointer and use pdfminer
#         file_obj.file.seek(0)
#         return extract_text_from_pdf_using_pdfminer(file_obj.file)                text += page.get_text()
#         if not text.strip():
#         raise ValueError("Empty text extracted with PyMuPDF.")
#             return text

# from pdfminer.high_level import extract_text        logging.warning(f"⚠️ PyMuPDF failed: {e}. Trying fallback with pdfminer.")

# def extract_text_from_pdf_using_pdfminer(file_path):
#     """Extracts text using pdfminer.six (fallback method)."""return extract_text_from_pdf_using_pdfminer(file_obj.file)
#     try:
#         return extract_text(file_path)
#     except Exception as e:
#         logging.error(f"pdfminer read error: {e}")
#         return Nonefrom pdfminer.high_level import extract_text

#     def extract_text_from_pdf_using_pdfminer(file_path):
# er.six (fallback method)."""
# def extract_text_from_docx(file):
#     """Extracts text from a DOCX file."""return extract_text(file_path)
#     try:
#         doc = docx.Document(file)
#         extracted_text = "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())
#         return extracted_text if extracted_text else None  # ✅ Return None instead of a fixed message
#     except Exception as e:
#         logging.error(f"DOCX file read error: {e}")
#         return Noneextract_text_from_docx(file):
#         """Extracts text from a DOCX file."""

# import asyncio = docx.Document(file)
# import uuidted_text = "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())
# import loggingurn None instead of a fixed message
# from langchain.text_splitter import RecursiveCharacterTextSplittert Exception as e:
# import re        logging.error(f"DOCX file read error: {e}")

# # async def store_text_in_chroma(text, source_name, model, user_id=None):
# #     try:
# #         safe_source = re.sub(r'\W+', '_', source_name)import asyncio

# #         existing_docs = chroma_collection.get(include=["documents", "metadatas"]) or {}
# #         existing_texts = existing_docs.get("documents", []) or []
# #         existing_metadata = existing_docs.get("metadatas", []) or []import re

# #         if user_id:del, user_id=None):
# #             for i, doc in enumerate(existing_texts):
# #                 if doc == text and existing_metadata[i].get("user_id") == user_id:
# #                     logging.info(f"🔄 Skipping duplicate for user {user_id} from {source_name}")
# #                     returning_docs = chroma_collection.get(include=["documents", "metadatas"]) or {}
# #         else:.get("documents", []) or []
# #             if text in existing_texts:
# #                 logging.info(f"🔄 Skipping duplicate entry from {source_name}")
# #                 return#         if user_id:

# #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)adata[i].get("user_id") == user_id:
# #         chunks = text_splitter.split_text(text)#                     logging.info(f"🔄 Skipping duplicate for user {user_id} from {source_name}")

# #         embeddings = await asyncio.to_thread(lambda: [model.encode(chunk).tolist() for chunk in chunks])
# #         ids = [f"{safe_source}_{uuid.uuid4()}" for _ in chunks]
# #         metadata = [{"source": source_name, "user_id": user_id} for _ in chunks]#                 logging.info(f"🔄 Skipping duplicate entry from {source_name}")

# #         # Not awaited — ChromaDB add is usually synchronous
# #         chroma_collection.add(siveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
# #             documents=chunks,it_text(text)
# #             embeddings=embeddings,
# #             ids=ids,ncio.to_thread(lambda: [model.encode(chunk).tolist() for chunk in chunks])
# #             metadatas=metadatads = [f"{safe_source}_{uuid.uuid4()}" for _ in chunks]
# #         )#         metadata = [{"source": source_name, "user_id": user_id} for _ in chunks]

# #         # metadata = [
# #         #    {"source": source_name, "user_id": user_id, "chunk_index": i}
# #         #    for i in range(len(chunks)) documents=chunks,
# #         # ]#             embeddings=embeddings,
#       ids=ids,
        
# #         # await asyncio.to_thread(
# #         #    chroma_collection.add,
# #         #    documents=chunks,
# #         #    embeddings=embeddings,": source_name, "user_id": user_id, "chunk_index": i}
# #         #    ids=ids,(chunks))
# #         #    metadatas=metadata
# #         # )   
        

# #         logging.info(f"✅ Stored {len(chunks)} chunks from {source_name} into ChromaDB for user {user_id}.")#         #    chroma_collection.add,
# nks,
# #     except Exception as e:
# #         logging.error(f"❌ Error storing in ChromaDB: {e}")#         #    ids=ids,
# data
# from datetime import datetime#         # )   

# def clean_metadata(value):
#     if value is None:"✅ Stored {len(chunks)} chunks from {source_name} into ChromaDB for user {user_id}.")
#         return "unknown"
#     if isinstance(value, (str, int, float, bool)):on as e:
#         return valueor(f"❌ Error storing in ChromaDB: {e}")
#     return str(value)
# from datetime import datetime

# import re, uuid, logging, asyncio
# from datetime import datetime
# from typing import Dict
# from fastapi import HTTPException, status, Querye, (str, int, float, bool)):
# from uuid import uuid4
# from langchain.text_splitter import RecursiveCharacterTextSplitter    return str(value)

# # --- Store text in ChromaDB --- #

# async def store_text_in_chroma(text: str, source_name: str, model, chroma_collection, user_id: str = None) -> Dict:
#     """Store text in ChromaDB with improved error handling and cleaned metadata."""ing import Dict
#     try:xception, status, Query
#         # Validate inputs
#         if not isinstance(text, str) or not text.strip():
#             logging.warning(f"⚠️ Invalid or empty text input from {source_name}")
#             return {"status": "error", "message": "Invalid or empty text input"}# --- Store text in ChromaDB --- #

#         if not hasattr(chroma_collection, 'add'): chroma_collection, user_id: str = None) -> Dict:
#             logging.error(f"❌ Invalid ChromaDB collection object")"""
#             return {"status": "error", "message": "Invalid ChromaDB collection"}    try:

#         safe_source = re.sub(r'\W+', '_', source_name)
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)ty text input from {source_name}")
#         chunks = text_splitter.split_text(text)            return {"status": "error", "message": "Invalid or empty text input"}

#         if not chunks:
#             logging.warning(f"⚠️ No valid chunks generated from text for {source_name}")
#             return {"status": "error", "message": "No valid chunks generated"}            return {"status": "error", "message": "Invalid ChromaDB collection"}

#         embeddings, valid_chunks = [], [](r'\W+', '_', source_name)
#         for chunk in chunks:itter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
#             try:
#                 embedding = model.encode(chunk).tolist()
#                 if embedding:
#                     embeddings.append(embedding)ks generated from text for {source_name}")
#                     valid_chunks.append(chunk)or", "message": "No valid chunks generated"}
#             except Exception as e:
#                 logging.warning(f"⚠️ Embedding error for chunk: {e}")d_chunks = [], []
#                 continue        for chunk in chunks:

#         if not embeddings:
#             logging.warning(f"⚠️ No valid embeddings generated for {source_name}")
#             return {"status": "error", "message": "No valid embeddings generated"}                    embeddings.append(embedding)

#             except Exception as e:
#                 logging.warning(f"⚠️ Embedding error for chunk: {e}")
#                 continue

#         if not embeddings:
#             logging.warning(f"⚠️ No valid embeddings generated for {source_name}")
#             return {"status": "error", "message": "No valid embeddings generated"}

#         ids = [f"{safe_source}_{uuid.uuid4()}" for _ in valid_chunks]
#         metadata = [{
#             "source": clean_metadata(source_name),
#             "user_id": clean_metadata(user_id),
#             "chunk_index": i,
#             "timestamp": datetime.utcnow().isoformat(),
#             "source_type": clean_metadata(source_name.split()[0] if source_name else "unknown"),
#             "query": clean_metadata(text[:100])
#         } for i in range(len(valid_chunks))]

#         # Use the actual ChromaDB collection object
#         chroma_collection.add(
#             documents=valid_chunks,
#             embeddings=embeddings,
#             ids=ids,
#             metadatas=metadata
#         )

#         logging.info(f"✅ Stored {len(valid_chunks)} chunks from '{source_name}' in ChromaDB")
#         return {"status": "success", "chunks_added": len(valid_chunks)}

#     except Exception as e:
#         logging.error(f"❌ ChromaDB storage error: {e}", exc_info=True)
#         return {"status": "error", "message": str(e)}



# # def print_all_chroma():
# #     """Prints all documents stored in ChromaDB."""
# #     try:
# #         all_data = chroma_collection.get(include=["documents"])
# #         if all_data["documents"]:
# #             for doc in all_data["documents"]:
# #                 print(f"Stored Document: {doc[:200]}...")  # Print first 200 chars
# #         else:
# #             print("⚠️ ChromaDB is empty!")
# #     except Exception as e:
# #         logging.error(f"Error retrieving ChromaDB data: {e}")



# def clear_chroma_collection(chroma_collection):
#     """Clears all data from the ChromaDB collection."""
#     try:
#         chroma_collection.delete(ids=chroma_collection.get()['ids'])
#         logging.info("ChromaDB collection cleared.")
#     except Exception as e:
#         logging.error(f"Error clearing ChromaDB: {e}")

# def print_all_chroma(chroma_collection):
#     """Prints all documents in the ChromaDB collection."""
#     try:
#         all_data = chroma_collection.get(include=["documents"])
#         if all_data["documents"]:
#     try:
#         # Create a thread pool for CPU-intensive operations
#         with ThreadPoolExecutor() as executor:
#             # Extract text based on file type
#             if file.content_type == "application/pdf":
#                 text = await asyncio.get_event_loop().run_in_executor(
#                     executor, extract_text_from_pdf, file
#                 )
#             elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                 text = await asyncio.get_event_loop().run_in_executor(
#                     executor, extract_text_from_docx, file
#                 )
#             else:
#                 text = await file.read()
#                 text = text.decode('utf-8')
                
#         # Split and process text in chunks
#         text_splitter = RecursiveCharacterTextSplitter(
#             chunk_size=1000,
#             chunk_overlap=200
#         )
#         chunks = text_splitter.split_text(text)
        
#         # Process chunks in parallel
#         tasks = [
#             store_text_in_chroma(chunk, f"File: {file.filename}", model)
#             for chunk in chunks
#         ]
#         await asyncio.gather(*tasks)
        
#         return f"Successfully processed {len(chunks)} chunks from {file.filename}"
        
#     except Exception as e:
#         logging.error(f"Error processing file: {e}")
#         return f"Error processing file: {str(e)}"

# async def return_extracted_data_from_chroma(query: str, model, chroma_collection, top_k=4):
#     """Returns relevant documents from ChromaDB based on similarity search."""
#     try:
#         query_embedding = await asyncio.to_thread(lambda: model.encode(query).tolist())
#         results = await asyncio.to_thread(
#             chroma_collection.query,
#             query_embeddings=[query_embedding],
#             n_results=top_k,
#             include=["documents", "metadatas"]
#         )

#         documents = results.get("documents", [[]])[0]
#         metadatas = results.get("metadatas", [[]])[0]

#         extracted_data = [
#             {
#                 "chunk": doc,
#                 "metadata": meta
#             }
#             for doc, meta in zip(documents, metadatas)
#         ]

#         logging.info(f"🔍 Retrieved {len(extracted_data)} chunks for query: {query}")
#         return extracted_data

#     except Exception as e:
#         logging.error(f"❌ Error retrieving from ChromaDB: {e}")
#         return []









import chromadb
from sentence_transformers import SentenceTransformer
import wikipedia
import requests
import os
import logging
import asyncio
from dotenv import load_dotenv
from bs4 import BeautifulSoup
from Backend.storage import get_chroma_collection
from langchain.text_splitter import RecursiveCharacterTextSplitter
from newspaper import Article
import PyPDF2
import docx
from docx import Document
import fitz
import httpx
import uuid
import validators
import re
import nltk
from nltk.tokenize import sent_tokenize
from transformers import pipeline
from pdfminer.high_level import extract_text as extract_text_from_pdf

# Load environment variables from .env file
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
GOOGLE_SEARCH_ENGINE_ID = os.getenv("GOOGLE_SEARCH_ENGINE_ID")
DUCKDUCKGO_API_URL = "https://api.duckduckgo.com/"
NEWS_API_KEY = os.getenv("NEWS_API_KEY")

logging.basicConfig(level=logging.DEBUG)

# Load embedding model
model = SentenceTransformer('all-MiniLM-L6-v2')
logging.info("✅ Model loaded successfully!")

# Load ChromaDB collection once
chroma_collection = get_chroma_collection()


# Initialize ChromaDB
chroma_client = chromadb.PersistentClient(path="./chroma_db")
collection = chroma_client.get_or_create_collection(name="study_docs")


nltk.download("punkt")

# Educational websites for scraping
EDUCATIONAL_SITES = [
    "https://www.khanacademy.org/",
    "https://www.coursera.org/",
    "https://www.edx.org/",
    "https://www.geeksforgeeks.org/"
]


# ✅ Load Summarization Model (One-time)
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")


async def search_google(query, model, retries=3):
    """Fetch search results from Google and fallback to DuckDuckGo on error."""
    url = "https://www.googleapis.com/customsearch/v1"
    params = {
        "q": query,
        "key": GOOGLE_API_KEY,
        "cx": GOOGLE_SEARCH_ENGINE_ID,
        "num": 5,
        "hl": "en"
    }

    for attempt in range(retries):
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(url, params=params, timeout=10)
                response.raise_for_status()
                data = response.json()

                if not data.get("items"):
                    logging.warning(f"⚠️ No results from Google. Switching to DuckDuckGo.")
                    return await search_duckduckgo(query, model)

                extracted_results = []
                for item in data["items"]:
                    title, snippet, link = item.get("title", ""), item.get("snippet", ""), item.get("link", "")

                    if not link or not validators.url(link) or any(x in link for x in ["youtube.com", "reddit.com"]):
                        continue

                    page_text = await scrape_and_summarize(link)
                    result_entry = f"🔹 **{title}**\n{snippet}\n🔗 {link}"
                    if page_text:
                        result_entry += f"\n📄 Summary: {page_text}"

                    extracted_results.append(result_entry)

                if extracted_results:
                    await store_text_in_chroma("\n".join(extracted_results), f"Google Data ({query})", model)

                return extracted_results

        except httpx.HTTPStatusError as e:
            if e.response.status_code == 429:
                wait_time = 2 ** attempt
                logging.warning(f"🚨 Rate limit hit! Waiting {wait_time}s...")
                await asyncio.sleep(wait_time)
            else:
                logging.error(f"❌ Google API error: {e}")
                break

        except Exception as e:
            logging.error(f"❌ Unexpected error: {e}")
            break
    
    if "ieeexplore.ieee.org" in url and response.status_code == 418:
       return {
        "type": "link-only",
        "message": (
            "🔒 The requested IEEE Xplore page cannot be accessed directly due to restrictions on automated scraping. "
            "You can view the document manually here:\n\n"
            f"{url}\n\n"
            "For metadata access, consider using the [IEEE Xplore API](https://developer.ieee.org/)."
        )
    }


    # Fallback if retries failed
    return await search_duckduckgo(query, model)


async def search_duckduckgo(query, model):
    """Fetches search results from DuckDuckGo if Google fails."""
    try:
        async with httpx.AsyncClient() as client:
            params = {"q": query, "format": "json"}
            response = await client.get(DUCKDUCKGO_API_URL, params=params, timeout=10)
            response.raise_for_status()
            data = response.json()

            if '50x.html' in str(response.url):
                logging.warning("⚠️ DuckDuckGo API returned a redirect. Possibly malformed input.")

            related_topics = data.get("RelatedTopics", [])
            extracted_results = []

            for topic in data.get("RelatedTopics", []):
              if "Text" in topic and "FirstURL" in topic:
                url = topic["FirstURL"]
                if not url.startswith("http"):
                  continue
                extracted_results.append(topic["Text"])

            if extracted_results:
                await store_text_in_chroma("\n".join(extracted_results), f"DuckDuckGo Data ({query})", model)

            if not extracted_results:
                logging.warning(f"⚠️ No useful DuckDuckGo results found for '{query}'")


            return extracted_results

    except Exception as e:
        logging.error(f"❌ DuckDuckGo API error: {e}")
        return None


async def fetch_latest_news(query):
    """Fetches the latest news from NewsAPI."""
    try:
        if not NEWS_API_KEY:
            raise ValueError("❌ Missing NewsAPI Key!")

        url = f"https://newsapi.org/v2/everything?q={query}&apiKey={NEWS_API_KEY}&language=en&pageSize=5"

        async with httpx.AsyncClient() as client:
            response = await client.get(url, timeout=10)
            response.raise_for_status()
            data = response.json()

        return [f"{article['title']}: {article['url']}" for article in data.get("articles", [])] if data.get("articles") else None
    except Exception as e:
        logging.error(f"❌ News API Error: {e}")
        return None



# import httpx
# from bs4 import BeautifulSoup
# import logging

# async def scrape_and_summarize(url):
#     """Scrapes a webpage and summarizes its content."""
#     if not url:
#         logging.error("❌ No URL found to scrape.")
#         return "No URL provided to scrape."

#     if not url.startswith("http://") and not url.startswith("https://"):
#         logging.error(f"❌ Invalid URL: {url}")
#         return "Invalid URL format."

#     try:
#         async with httpx.AsyncClient(timeout=5.0) as client:  # shorter timeout
#             response = await client.get(url)
#             response.raise_for_status()

#         soup = BeautifulSoup(response.text, "html.parser")

#         # Remove unnecessary tags
#         for tag in soup(["script", "style", "noscript", "header", "footer", "nav", "aside"]):
#             tag.extract()

#         # Extract meaningful content
#         content_elements = soup.find_all(["p", "div", "article", "section"])
#         text = " ".join([elem.get_text(separator=" ", strip=True) for elem in content_elements]).strip()

#         if len(text) < 100:
#             return None  # too little content
#         elif len(text) <= 500:
#             return text  # already concise enough

#         # Truncate and summarize
#         text = text[:1024]
#         try:
#             summary = summarizer(text, max_length=150, min_length=50, do_sample=False)[0]['summary_text']
#             return summary
#         except Exception as e:
#             logging.error(f"❌ Summarization error: {e}")
#             return text[:300]  # fallback to raw excerpt

#     except httpx.HTTPStatusError as e:
#         logging.error(f"❌ HTTP error during scraping {url}: {e}")
#     except httpx.RequestError as e:
#         logging.error(f"❌ Request failed for {url}: {e}")
#     except Exception as e:
#         logging.error(f"❌ General scraping error for {url}: {e}")

#     return None  # final fallback if everything fails

        



import httpx
import logging
from bs4 import BeautifulSoup
from urllib.parse import urlparse
from transformers import pipeline

# Initialize only once (ensure you do this outside if you're importing it)
summarizer = pipeline("summarization")

def is_valid_url(url: str) -> bool:
    """Validates whether a string is a proper URL."""
    try:
        parsed = urlparse(url)
        return parsed.scheme in ("http", "https") and parsed.netloc != ""
    except Exception:
        return False

async def scrape_and_summarize(url_or_query: str):
    """Scrapes a webpage if input is a valid URL and summarizes it. Otherwise returns None."""
    
    if not url_or_query:
        logging.error("❌ No input received for scraping.")
        return None

    # Not a valid URL — skip scraping
    if not is_valid_url(url_or_query):
        logging.info(f"⚠️ Input is not a URL. Skipping scraping: {url_or_query}")
        return None

    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            response = await client.get(url_or_query)
            response.raise_for_status()

        soup = BeautifulSoup(response.text, "html.parser")

        # Remove non-content elements
        for tag in soup(["script", "style", "noscript", "header", "footer", "nav", "aside"]):
            tag.extract()

        # Extract text
        content_elements = soup.find_all(["p", "div", "article", "section"])
        text = " ".join([elem.get_text(separator=" ", strip=True) for elem in content_elements]).strip()

        if len(text) < 100:
            logging.warning("⚠️ Scraped content too short.")
            return None
        elif len(text) <= 500:
            logging.info("ℹ️ Short content. Skipping summarization.")
            return text

        # Truncate for summarization
        text = text[:1024]

        try:
            summary = summarizer(text, max_length=150, min_length=50, do_sample=False)[0]['summary_text']
            logging.info("✅ Summary successfully generated.")
            return summary
        except Exception as e:
            logging.error(f"❌ Summarization error: {e}")
            return text[:300]  # fallback

    except httpx.HTTPStatusError as e:
        logging.error(f"❌ HTTP error while scraping {url_or_query}: {e}")
    except httpx.RequestError as e:
        logging.error(f"❌ Request failed for {url_or_query}: {e}")
    except Exception as e:
        logging.error(f"❌ General scraping error: {e}")

    return None


# ✅ Helper Function to Extract First Full Sentences (Avoid Mid-Cutoff)
def extract_first_sentences(text, max_length=500):
    """Extracts the first few full sentences without cutting off mid-way."""
    sentences = text.split(". ")
    output = ""
    for sentence in sentences:
        if len(output) + len(sentence) > max_length:
            break
        output += sentence + ". "
    return output.strip()


# def search_web(query):
#     """Performs a web search using DuckDuckGo API."""
#     try:
#         url = f"https://api.duckduckgo.com/?q={query}&format=json"
#         response = requests.get(url, timeout=10)
#         response.raise_for_status()
#         data = response.json()
#         results = [topic["Text"] for topic in data.get("RelatedTopics", []) if "Text" in topic]
#         return results if results else ["No relevant results found."]
#     except requests.exceptions.RequestException as e:
#         logging.error(f"Web search failed: {e}")
#         return ["Web search error."]
#     except Exception as e:
#         logging.error(f"Unexpected web search error: {e}")
#         return ["Unexpected error occurred during web search."]

import requests


def search_web(query):
    """Performs a web search using DuckDuckGo API."""
    try:
        url = f"https://api.duckduckgo.com/?q={query}&format=json"
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        data = response.json()

        results = [topic["Text"] for topic in data.get("RelatedTopics", []) if "Text" in topic]
        return results if results else ["No relevant results found."]
    except requests.exceptions.RequestException as e:
        logging.error(f"Web search failed: {e}")
        return ["Web search error."]


async def scrape_educational_websites(query, model):
    """Scrapes educational websites for study material using semantic similarity."""
    results = []
    query_embedding = model.encode(query)

    for site in EDUCATIONAL_SITES:
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(site, timeout=10)
                response.raise_for_status()

            soup = BeautifulSoup(response.text, "html.parser")
            content_elements = soup.find_all(['p', 'div', 'article', 'section', 'span'])
            text = " ".join([element.get_text(separator=" ", strip=True) for element in content_elements])

            text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
            chunks = text_splitter.split_text(text)

            for chunk in chunks:
                chunk_embedding = model.encode(chunk)
                similarity = chunk_embedding @ query_embedding
                if similarity > 0.6:
                    results.append((similarity, f"From {site}: {chunk}"))

        except Exception as e:
            logging.error(f"Error fetching or processing {site}: {e}")

    results.sort(key=lambda x: x[0], reverse=True)
    return [result[1] for result in results[:3]]


import wikipedia

def async_wikipedia_search(query):
    """Fetch Wikipedia summary with proper error handling and debugging."""
    try:
        print(f"🔎 Searching Wikipedia for: {query}")
        search_results = wikipedia.search(query)

        print(f"📜 Wikipedia Search Results: {search_results}")

        if search_results:
           page = wikipedia.page(search_results[0])
           content = page.content
        else:
           logging.warning("⚠️ No relevant Wikipedia result found for '%s'", query)

        if not search_results:
            return None

        for result in search_results[:3]:
            try:
                page = wikipedia.page(result, auto_suggest=True)
                summary = wikipedia.summary(page.title, sentences=3)

                if query.lower() in page.title.lower() and len(summary) > 50:
                    return summary

            except wikipedia.exceptions.DisambiguationError as e:
                print(f"⚠️ Disambiguation Error: {e}")
                for option in e.options[:3]:
                    try:
                        summary = wikipedia.summary(option, sentences=3)
                        if len(summary) > 50:
                            return summary
                    except:
                        continue

            except wikipedia.exceptions.PageError:
                print(f"❌ PageError for {result}")
                continue

        return None

    except Exception as e:
        print(f"❌ Wikipedia API General Error: {e}")
        return None

    

async def async_scrape_web(query):
    """Asynchronously scrape educational websites for content if needed."""
    return scrape_educational_websites(query, model) or None



async def retrieve_relevant_text(query: str, model, top_k=3):
    """Retrieve relevant study material asynchronously from ChromaDB."""
    try:
        query_embedding = model.encode(query).tolist()
        available_docs = chroma_collection.count()
        num_results = max(1, min(top_k, available_docs))

        results = await asyncio.to_thread(
            chroma_collection.query, 
            query_embeddings=[query_embedding], 
            n_results=num_results
        )

        if not results or "documents" not in results or not results["documents"]:
            logging.warning("⚠️ No relevant documents found in ChromaDB.")
            return "No relevant results found."

        retrieved_docs = results["documents"][0] if results["documents"] else []
        distances = results.get("distances", [[]])[0] if results.get("distances") else []

        cleaned_results = []
        for idx, text in enumerate(retrieved_docs):
            if text and distances[idx] < 0.3:
                text = re.sub(r"[\n\t]+", " ", text)  # Remove line breaks/tabs
                text = re.sub(r"http\S+", "", text)  # Remove links
                text = re.sub(r"[^a-zA-Z0-9.,!?'\- ]+", "", text)  # ✅ Keeps hyphens and apostrophes
                text = re.sub(r"\s+", " ", text).strip()
                
                if len(text.split()) > 5:
                    cleaned_results.append(text)

        return "\n\n".join(cleaned_results) if cleaned_results else "No relevant results found."

    except Exception as e:
        logging.error(f"❌ ChromaDB retrieval error: {e}")
        return f"Error retrieving from ChromaDB: {str(e)}"
    


import fitz  # PyMuPDF

def extract_text_from_pdf(file_obj):
    """Attempts text extraction from PDF using PyMuPDF, with pdfminer fallback."""
    try:
        with fitz.open(stream=file_obj.file.read(), filetype="pdf") as doc:
            text = ""
            for page in doc:
                text += page.get_text()
        if not text.strip():
            raise ValueError("Empty text extracted with PyMuPDF.")
        return text
    except Exception as e:
        logging.warning(f"⚠️ PyMuPDF failed: {e}. Trying fallback with pdfminer.")
        # Rewind file pointer and use pdfminer
        file_obj.file.seek(0)
        return extract_text_from_pdf_using_pdfminer(file_obj.file)


    

from pdfminer.high_level import extract_text

def extract_text_from_pdf_using_pdfminer(file_path):
    """Extracts text using pdfminer.six (fallback method)."""
    try:
        return extract_text(file_path)
    except Exception as e:
        logging.error(f"pdfminer read error: {e}")
        return None

    

def extract_text_from_docx(file):
    """Extracts text from a DOCX file."""
    try:
        doc = docx.Document(file)
        extracted_text = "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())
        return extracted_text if extracted_text else None  # ✅ Return None instead of a fixed message
    except Exception as e:
        logging.error(f"DOCX file read error: {e}")
        return None
    

import asyncio
import uuid
import logging
from langchain.text_splitter import RecursiveCharacterTextSplitter
import re

# async def store_text_in_chroma(text, source_name, model, user_id=None):
#     try:
#         safe_source = re.sub(r'\W+', '_', source_name)

#         existing_docs = chroma_collection.get(include=["documents", "metadatas"]) or {}
#         existing_texts = existing_docs.get("documents", []) or []
#         existing_metadata = existing_docs.get("metadatas", []) or []

#         if user_id:
#             for i, doc in enumerate(existing_texts):
#                 if doc == text and existing_metadata[i].get("user_id") == user_id:
#                     logging.info(f"🔄 Skipping duplicate for user {user_id} from {source_name}")
#                     return
#         else:
#             if text in existing_texts:
#                 logging.info(f"🔄 Skipping duplicate entry from {source_name}")
#                 return

#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
#         chunks = text_splitter.split_text(text)

#         embeddings = await asyncio.to_thread(lambda: [model.encode(chunk).tolist() for chunk in chunks])
#         ids = [f"{safe_source}_{uuid.uuid4()}" for _ in chunks]
#         metadata = [{"source": source_name, "user_id": user_id} for _ in chunks]

#         # Not awaited — ChromaDB add is usually synchronous
#         chroma_collection.add(
#             documents=chunks,
#             embeddings=embeddings,
#             ids=ids,
#             metadatas=metadata
#         )

#         # metadata = [
#         #    {"source": source_name, "user_id": user_id, "chunk_index": i}
#         #    for i in range(len(chunks))
#         # ]

        
#         # await asyncio.to_thread(
#         #    chroma_collection.add,
#         #    documents=chunks,
#         #    embeddings=embeddings,
#         #    ids=ids,
#         #    metadatas=metadata
#         # )   


#         logging.info(f"✅ Stored {len(chunks)} chunks from {source_name} into ChromaDB for user {user_id}.")

#     except Exception as e:
#         logging.error(f"❌ Error storing in ChromaDB: {e}")

from datetime import datetime

async def store_text_in_chroma(text, source_name, model, user_id=None):
    try:
        safe_source = re.sub(r'\W+', '_', source_name)

        existing_docs = chroma_collection.get(include=["documents", "metadatas"]) or {}
        existing_texts = existing_docs.get("documents", []) or []
        existing_metadata = existing_docs.get("metadatas", []) or []

        if user_id:
            for i, doc in enumerate(existing_texts):
                if doc == text and existing_metadata[i].get("user_id") == user_id:
                    logging.info(f"🔄 Skipping duplicate for user {user_id} from {source_name}")
                    return {"status": "skipped", "reason": "duplicate"}
        else:
            if text in existing_texts:
                logging.info(f"🔄 Skipping duplicate entry from {source_name}")
                return {"status": "skipped", "reason": "duplicate"}

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
        chunks = text_splitter.split_text(text)

        embeddings = await asyncio.to_thread(lambda: [model.encode(chunk).tolist() for chunk in chunks])
        ids = [f"{safe_source}_{uuid.uuid4()}" for _ in chunks]
        metadata = [
          {
             "source": source_name,
             "user_id": user_id,
             "chunk_index": i,
             "timestamp": datetime.utcnow().isoformat(),
             "source_type": source_name.split()[0],  # e.g., "Google", "DuckDuckGo"
             "query": text[:100],  # Optional: for trace/debug
        # "source_url": "https://example.com",  # if available from search
        # "confidence": 0.9  # if using a scoring system
          }
          for i in range(len(chunks))
        ]

        await asyncio.to_thread(
            chroma_collection.add,
            documents=chunks,
            embeddings=embeddings,
            ids=ids,
            metadatas=metadata
        )

        logging.info(f"✅ Stored {len(chunks)} chunks from {source_name} into ChromaDB for user {user_id}.")
        logging.debug(f"🧩 First chunk preview: {chunks[0][:100]}")

        return {"status": "success", "chunks_added": len(chunks)}

    except Exception as e:
        logging.error(f"❌ Error storing in ChromaDB: {e}")
        return {"status": "error", "message": str(e)}


# def print_all_chroma():
#     """Prints all documents stored in ChromaDB."""
#     try:
#         all_data = chroma_collection.get(include=["documents"])
#         if all_data["documents"]:
#             for doc in all_data["documents"]:
#                 print(f"Stored Document: {doc[:200]}...")  # Print first 200 chars
#         else:
#             print("⚠️ ChromaDB is empty!")
#     except Exception as e:
#         logging.error(f"Error retrieving ChromaDB data: {e}")



def clear_chroma_collection(chroma_collection):
    """Clears all data from the ChromaDB collection."""
    try:
        chroma_collection.delete(ids=chroma_collection.get()['ids'])
        logging.info("ChromaDB collection cleared.")
    except Exception as e:
        logging.error(f"Error clearing ChromaDB: {e}")

def print_all_chroma(chroma_collection):
    """Prints all documents in the ChromaDB collection."""
    try:
        all_data = chroma_collection.get(include=["documents"])
        if all_data["documents"]:
            for doc in all_data["documents"]:
                print(f"Document: {doc}")
        else:
            print("ChromaDB collection is empty.")
    except Exception as e:
        logging.error(f"Error printing ChromaDB data: {e}")

async def process_uploaded_file(file, model, chroma_collection):
    """Processes uploaded files and stores embeddings in ChromaDB."""
    try:
        logging.info(f"📁 Received file: {file.filename}")

        # Step 1: Extract text
        if file.filename.endswith(".pdf"):
            logging.info("📄 Detected PDF format.")
            text = extract_text_from_pdf(file)
        elif file.filename.endswith(".docx"):
            logging.info("📄 Detected DOCX format.")
            text = extract_text_from_docx(file)
        else:
            return "Unsupported file format. Please upload a PDF or DOCX file."

        logging.info(f"📝 Extracted text length: {len(text) if text else 0}")

        # Step 2: Store in Chroma
        if text:
            await store_text_in_chroma(text, file.filename, model, chroma_collection)
            return "File processed and stored successfully."
        
        return "No valid text found in the uploaded file."

    except Exception as e:
        logging.error(f"❌ Error processing uploaded file: {e}", exc_info=True)
        return "Error processing uploaded file."

async def return_extracted_data_from_chroma(query: str, model, chroma_collection, top_k=4):
    """Returns relevant documents from ChromaDB based on similarity search."""
    try:
        query_embedding = await asyncio.to_thread(lambda: model.encode(query).tolist())
        results = await asyncio.to_thread(
            chroma_collection.query,
            query_embeddings=[query_embedding],
            n_results=top_k,
            include=["documents", "metadatas"]
        )

        documents = results.get("documents", [[]])[0]
        metadatas = results.get("metadatas", [[]])[0]

        extracted_data = [
            {
                "chunk": doc,
                "metadata": meta
            }
            for doc, meta in zip(documents, metadatas)
        ]

        logging.info(f"🔍 Retrieved {len(extracted_data)} chunks for query: {query}")
        return extracted_data

    except Exception as e:
        logging.error(f"❌ Error retrieving from ChromaDB: {e}")
        return []
