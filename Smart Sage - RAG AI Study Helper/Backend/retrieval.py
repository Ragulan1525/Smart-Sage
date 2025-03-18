# import chromadb
# from sentence_transformers import SentenceTransformer
# import wikipedia
# import requests
# from bs4 import BeautifulSoup
# import PyPDF2
# import docx
# from Backend.storage import get_chroma_collection
# import logging

# logging.basicConfig(level=logging.DEBUG)

# # Load embedding model
# model = SentenceTransformer('all-MiniLM-L6-v2')

# print("Model loaded successfully!")


# # Educational websites for scraping
# EDUCATIONAL_SITES = [
#     "https://www.khanacademy.org/",
#     "https://www.coursera.org/",
#     "https://www.edx.org/"
# ]

# def scrape_educational_websites(query):
#     """Scrapes educational websites for study material."""
#     for site in EDUCATIONAL_SITES:
#         try:
#             response = requests.get(site, timeout=5)
#             response.raise_for_status()  # Raise HTTPError for bad responses (4xx and 5xx)
#             soup = BeautifulSoup(response.text, "html.parser")
#             paragraphs = [p.text for p in soup.find_all("p")]
            
#             # Search for the query in the extracted text
#             for para in paragraphs:
#                 if query.lower() in para.lower():
#                     return f"From {site}: {para[:300]}..."  # Return relevant snippet

#         except requests.exceptions.RequestException as e:
#             logging.error(f"Error fetching {site}: {e}")
#             continue
#     return None  # If no relevant content is found


# def retrieve_relevant_text(query):
#     """Retrieves educational content from ChromaDB, Wikipedia, or Web Scraping."""
#     logging.info(f"🔍 Searching for query: {query}")

#     # 🔹 1️⃣ Check ChromaDB first
#     try:
#         collection = get_chroma_collection()
#         query_embedding = model.encode(query).tolist()
#         results = collection.query(query_embeddings=[query_embedding], n_results=3)

#         retrieved_docs = results.get("documents", [])
#         if retrieved_docs and retrieved_docs[0]:
#             logging.info(f"✅ Retrieved from ChromaDB: {retrieved_docs[0]}")
#             return "ChromaDB Result: " + " ".join(retrieved_docs[0])[:500]
#         else:
#             logging.warning("⚠️ No relevant results in ChromaDB.")
#     except Exception as e:
#         logging.error(f"❌ ChromaDB Retrieval Error: {e}")

#     # 🔹 2️⃣ Check Wikipedia if ChromaDB fails
#     try:
#         search_results = wikipedia.search(query)
#         if search_results:
#             wiki_summary = wikipedia.summary(search_results[0], sentences=2)
#             logging.info(f"✅ Wikipedia Result: {wiki_summary}")
#             return f"Wikipedia says: {wiki_summary}"
#     except Exception as e:
#         logging.error(f"❌ Wikipedia error: {e}")

#     # 🔹 3️⃣ Try Web Scraping if Wikipedia fails
#     web_result = scrape_educational_websites(query)
#     if web_result:
#         logging.info(f"✅ Web Scraping Result: {web_result}")
#         return web_result

#     # 🔹 4️⃣ As a LAST RESORT, return extracted text from PDFs
#     return "No relevant data found in ChromaDB or the web."





# #%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%5

# def search_web(query):
#     """Performs a web search using DuckDuckGo API."""
#     try:
#         url = f"https://api.duckduckgo.com/?q={query}&format=json"
#         response = requests.get(url, timeout=5)
#         response.raise_for_status()

#         data = response.json()
#         results = [topic["Text"] for topic in data.get("RelatedTopics", []) if "Text" in topic]

#         return results if results else ["No relevant results found."]
#     except Exception as e:
#         logging.error(f"Web search failed: {e}", exc_info=True)
#         return ["Error occurred while searching."]


# def extract_text_from_pdf(file):
#     """Extracts text from a PDF file."""
#     try:
#         reader = PyPDF2.PdfReader(file)
#         text = []
        
#         for page in reader.pages:
#             extracted = page.extract_text()
#             if extracted:
#                 text.append(extracted.strip())

#         full_text = "\n".join(text).strip()
#         return full_text if full_text else "No readable text found in the PDF."

#     except Exception as e:
#         logging.error(f"PDF extraction error: {e}")
#         return "Error reading PDF file."

    
# def extract_text_from_docx(file):
#     """Extracts text from a DOCX file."""
#     try:
#         doc = docx.Document(file)
#         text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]  # Corrected line
#         return "\n".join(text).strip() or "No readable text found in the DOCX file."
#     except Exception as e:
#         logging.error(f"DOCX extraction error: {e}")
#         return "Error reading DOCX file."
    

# def store_text_in_chroma(text):
#     """Stores extracted text into ChromaDB as small chunks."""
#     try:
#         collection = get_chroma_collection()
        
#         # 🔹 Break into 500-character chunks
#         chunks = [text[i:i+500] for i in range(0, len(text), 500)]  

#         embeddings = [model.encode(chunk).tolist() for chunk in chunks]

#         # 🔹 Store text & embeddings
#         collection.add(
#             documents=chunks,
#             embeddings=embeddings,
#             ids=[str(hash(chunk)) for chunk in chunks]  
#         )
#         logging.info(f"✅ Stored {len(chunks)} chunks in ChromaDB!")
#     except Exception as e:
#         logging.error(f"❌ Error storing data in ChromaDB: {e}")




# def process_uploaded_file(file):
#     """Processes uploaded PDF, DOCX, or TXT files and stores embeddings in ChromaDB."""
#     try:
#         if file.filename.endswith(".pdf"):
#             text = extract_text_from_pdf(file.file)
#         elif file.filename.endswith(".docx"):
#             text = extract_text_from_docx(file.file)
#         elif file.filename.endswith(".txt"):
#             text = file.file.read().decode("utf-8").strip()
#         else:
#             raise ValueError("Unsupported file format. Please upload PDF, DOCX, or TXT.")

#         if text:
#             store_text_in_chroma(text)  # Store extracted text in ChromaDB
#             return "File processed and stored successfully."

#         return "No valid text found in the uploaded file."
#     except Exception as e:
#         logging.error(f"File processing error: {e}")
#         return "Error processing file."



#%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%55555

import chromadb
from sentence_transformers import SentenceTransformer
import wikipedia
import requests
from bs4 import BeautifulSoup
import PyPDF2
import docx
from Backend.storage import get_chroma_collection
import logging
from langchain.text_splitter import RecursiveCharacterTextSplitter
from newspaper import Article

logging.basicConfig(level=logging.DEBUG)

# Load embedding model
model = SentenceTransformer('all-MiniLM-L6-v2')
logging.info("Model loaded successfully!")

# Educational websites for scraping
EDUCATIONAL_SITES = [
    "https://www.khanacademy.org/",
    "https://www.coursera.org/",
    "https://www.edx.org/",
    "https://www.geeksforgeeks.org/"
]

def scrape_educational_websites(query, model):
    """Scrapes educational websites for study material using semantic similarity."""
    results = []
    query_embedding = model.encode(query)

    for site in EDUCATIONAL_SITES:
        try:
            article = Article(site)
            article.download()
            article.parse()
            text = article.text

            text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
            chunks = text_splitter.split_text(text)

            for chunk in chunks:
                chunk_embedding = model.encode(chunk)
                similarity = chunk_embedding @ query_embedding  # Dot product for cosine similarity
                if similarity > 0.6:  # Adjust threshold as needed
                    results.append((similarity, f"From {site}: {chunk}"))

        except Exception as e:
            logging.error(f"Error fetching or processing {site}: {e}")

    results.sort(key=lambda x: x[0], reverse=True)  # Sort by similarity
    return [result[1] for result in results[:3]]  # Return top 3 results

def retrieve_relevant_text(query: str, model):
    """Retrieves educational content from ChromaDB, Wikipedia, or Web Scraping."""
    logging.info(f"🔍 Searching for query: {query}")
    
    try:
        client = chromadb.PersistentClient()
        collection = client.get_collection("study_data")
        query_embedding = model.encode(query).tolist()
        results = collection.query(query_embeddings=[query_embedding], n_results=3)
        retrieved_docs = results.get("documents", [])
        
        if retrieved_docs and retrieved_docs[0]:
            logging.info(f"✅ Retrieved from ChromaDB: {retrieved_docs[0]}")
            return "ChromaDB Result: " + " ".join(retrieved_docs[0])
        else:
            logging.warning("⚠️ No relevant results in ChromaDB.")
    except Exception as e:
        logging.error(f"❌ ChromaDB Retrieval Error: {e}")
    
    # Check Wikipedia if ChromaDB fails
    try:
        search_results = wikipedia.search(query)
        if search_results:
            wiki_summary = wikipedia.summary(search_results[0], sentences=2)  # Limit to 2 sentences for speed
            logging.info(f"✅ Wikipedia Result: {wiki_summary}")
            return f"Wikipedia says: {wiki_summary}"
    except (wikipedia.exceptions.PageError, wikipedia.exceptions.DisambiguationError, Exception) as e:
        logging.error(f"❌ Wikipedia Error: {e}", exc_info=True)
    
    # Try Web Scraping if Wikipedia fails
    web_results = scrape_educational_websites(query, model)
    if web_results:
        logging.info(f"✅ Web Scraping Result: {web_results}")
        return "\n".join(web_results)
    
    return "No relevant data found in ChromaDB or the web."


def search_web(query):
    """Performs a web search using DuckDuckGo API."""
    try:
        url = f"https://api.duckduckgo.com/?q={query}&format=json"
        response = requests.get(url, timeout=10) #increased timeout.
        response.raise_for_status()

        data = response.json()
        results = [topic["Text"] for topic in data.get("RelatedTopics", []) if "Text" in topic]

        return results if results else ["No relevant results found."]
    except requests.exceptions.RequestException as e:
        logging.error(f"Web search failed: {e}")
        return ["Web search error."]
    except Exception as e:
        logging.error(f"Unexpected web search error: {e}")
        return ["Unexpected error occurred during web search."]

def extract_text_from_pdf(file):
    """Extracts text from a PDF file."""
    try:
        reader = PyPDF2.PdfReader(file)
        text = []
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text.append(extracted.strip())
        full_text = "\n".join(text).strip()
        return full_text if full_text else "No readable text found in the PDF."
    except PyPDF2.errors.PdfReadError as e:
        logging.error(f"PDF read error: {e}")
        return "Error reading PDF file."
    except Exception as e:
        logging.error(f"Unexpected PDF extraction error: {e}")
        return "Unexpected error reading PDF file."

def extract_text_from_docx(file):
    """Extracts text from a DOCX file."""
    try:
        doc = docx.Document(file)
        text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        return "\n".join(text).strip() or "No readable text found in the DOCX file."
    except docx.opc.exceptions.PackageNotFoundError as e:
        logging.error(f"DOCX file not found: {e}")
        return "Error reading DOCX file."
    except Exception as e:
        logging.error(f"Unexpected DOCX extraction error: {e}")
        return "Unexpected error reading DOCX file."

def store_text_in_chroma(text, model):
    """Stores extracted text into ChromaDB using semantic chunking."""
    try:
        collection = get_chroma_collection()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
        chunks = text_splitter.split_text(text)
        embeddings = [model.encode(chunk).tolist() for chunk in chunks]
        collection.add(
            documents=chunks,
            embeddings=embeddings,
            ids=[str(hash(chunk)) for chunk in chunks]
        )
        logging.info(f"✅ Stored {len(chunks)} chunks in ChromaDB!")
    except Exception as e:
        logging.error(f"❌ Error storing data in ChromaDB: {e}")

def process_uploaded_file(file, model):
    """Processes uploaded files and stores embeddings in ChromaDB."""
    try:
        if file.filename.endswith(".pdf"):
            text = extract_text_from_pdf(file.file)
        elif file.filename.endswith(".docx"):
            text = extract_text_from_docx(file.file)
        elif file.filename.endswith(".txt"):
            text = file.file.read().decode("utf-8").strip()
        else:
            raise ValueError("Unsupported file format. Please upload PDF, DOCX, or TXT.")

        if text:
            store_text_in_chroma(text, model)
            return "File processed and stored successfully."
        return "No valid text found in the uploaded file."
    except ValueError as ve:
        logging.error(f"File processing error: {ve}")
        return f"File processing error: {ve}"
    except Exception as e:
        logging.error(f"Unexpected file processing error: {e}")
        return "Unexpected error processing file."